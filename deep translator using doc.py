import streamlit as st
from docx import Document
from deep_translator import GoogleTranslator

# Set up translator
translator = GoogleTranslator(source="en", target="ur")

# Streamlit interface
st.title("Document Translation to Urdu")
st.write("Upload a Word document (.docx) in English to translate it into Urdu.")

# File upload option
uploaded_file = st.file_uploader("Upload your Word file", type="docx")

if uploaded_file:
    try:
        # Open the uploaded Word document
        doc = Document(uploaded_file)

        # Function to translate text using Deep Translator
        def translate_to_urdu(text):
            try:
                return translator.translate(text)
            except Exception as e:
                st.error(f"Translation error: {e}")
                return text

        # Process each paragraph in the document
        total_paragraphs = len(doc.paragraphs)
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Only translate non-empty paragraphs
                original_text = paragraph.text
                translated_text = translate_to_urdu(original_text)
                paragraph.text = translated_text
            print(f"Paragraph {i + 1} of {total_paragraphs} translated")

        # Save the translated document
        output_file = "translated-urdu.docx"
        doc.save(output_file)

        st.success("Translation completed!")
        with open(output_file, "rb") as f:
            st.download_button(
                label="Download Translated Document",
                data=f,
                file_name="translated-urdu.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        st.error(f"An error occurred: {e}")
